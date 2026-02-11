from docx import Document
import os

path = r"c:\Users\kingj\Documents\Programming3_2\Project\Demo_final\demo_fipro\docs\แบบประเมิน\วท.ป.1 แบบประเมินกรอบของโครงงาน.docx"
if os.path.exists(path):
    doc = Document(path)
    print("--- Paragraphs ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    print("--- Tables ---")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(cell.text.replace('\n', ' '))
else:
    print("File not found")
